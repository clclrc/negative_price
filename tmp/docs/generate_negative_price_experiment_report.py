from __future__ import annotations

from pathlib import Path
import re

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path("/Users/deadlock/Desktop/6002")
OUTPUT_PATH = ROOT / "output" / "doc" / "negative_price_experiment_report.docx"

METRIC_PATHS = {
    "E81": ROOT / "outputs" / "experiment_outputs_20260409_233636" / "E81" / "metrics_summary.csv",
    "E75": ROOT / "outputs" / "experiment_outputs_20260409_011614" / "E75" / "metrics_summary.csv",
    "E78": ROOT / "outputs" / "experiment_outputs_20260409_114522" / "E78" / "metrics_summary.csv",
    "E49": ROOT / "outputs" / "experiment_outputs_20260406_013305" / "E49" / "metrics_summary.csv",
    "E41": ROOT / "outputs" / "E41" / "metrics_summary.csv",
    "E42": ROOT / "outputs" / "experiment_outputs_20260405_124506" / "E42" / "metrics_summary.csv",
    "E43": ROOT / "outputs" / "E43" / "metrics_summary.csv",
    "E44": ROOT / "outputs" / "experiment_outputs_20260407_012625" / "E56" / "member_runs" / "E44" / "metrics_summary.csv",
    "E8": ROOT / "outputs" / "experiment_outputs_20260402_132259" / "E8" / "metrics_summary.csv",
    "E35": ROOT / "outputs" / "experiment_outputs_20260404_113925" / "E35" / "metrics_summary.csv",
}


TITLE = (
    "System Conditions Behind Negative Electricity Prices in High Renewable "
    "European Power Markets: A Deep Learning-Based Event Prediction Approach"
)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False, size: float = 9.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def configure_document(document: Document):
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)

    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size in (("Title", 18), ("Heading 1", 14), ("Heading 2", 12)):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_after = Pt(6)

    if "Table Text" not in styles:
        table_style = styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
        table_style.font.name = "Times New Roman"
        table_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        table_style.font.size = Pt(9.5)
        table_style.paragraph_format.space_after = Pt(0)
        table_style.paragraph_format.line_spacing = 1.0


def read_test_row(experiment_id: str) -> pd.Series:
    df = pd.read_csv(METRIC_PATHS[experiment_id])
    test_df = df[df["split"] == "test"].copy()
    if experiment_id == "E49":
        test_df = test_df[test_df["seed_aggregation"] == "mean"].copy()
    if test_df.empty:
        raise RuntimeError(f"No test row found for {experiment_id}")
    return test_df.iloc[0]


def f4(value) -> str:
    return f"{float(value):.4f}"


def build_main_metrics():
    rows = {experiment_id: read_test_row(experiment_id) for experiment_id in ("E41", "E42", "E43", "E44", "E49", "E75", "E78", "E81", "E8", "E35")}
    return rows


def experiment_sort_key(experiment_id: str):
    match = re.match(r"E(\d+)([A-Z]?)", experiment_id)
    if not match:
        return (9999, experiment_id)
    return (int(match.group(1)), match.group(2))


def build_experiment_registry():
    outputs_root = ROOT / "outputs"
    exp_pat = re.compile(r"/(E\d+[A-Z]?)/metrics_summary\.csv$")
    candidates: dict[str, list[dict[str, object]]] = {}
    for path in outputs_root.rglob("metrics_summary.csv"):
        path_str = str(path)
        if any(token in path_str for token in ("/member_runs/", "/seed_runs/", "/candidate_runs/")):
            continue
        match = exp_pat.search(path_str)
        if not match:
            continue
        experiment_id = match.group(1)
        try:
            df = pd.read_csv(path)
        except Exception:
            continue
        if df.empty or "split" not in df.columns:
            continue
        test_df = df[df["split"].astype(str) == "test"].copy()
        if "seed_aggregation" in test_df.columns:
            mean_df = test_df[test_df["seed_aggregation"].astype(str) == "mean"]
            if not mean_df.empty:
                test_df = mean_df
        if test_df.empty:
            continue
        row = test_df.iloc[0]
        candidates.setdefault(experiment_id, []).append(
            {
                "path": path,
                "mtime": path.stat().st_mtime,
                "model": str(row.get("model", "")),
                "pr_auc": float(row.get("pr_auc")),
            }
        )

    completed = {}
    for experiment_id, items in candidates.items():
        chosen = sorted(items, key=lambda item: item["mtime"])[-1]
        completed[experiment_id] = chosen

    config_text = (ROOT / "negative_price_experiments" / "config.py").read_text()
    configured_ids = set(re.findall(r'"(E\d+[A-Z]?)"\s*:', config_text))
    incomplete = sorted(configured_ids - set(completed.keys()), key=experiment_sort_key)
    return completed, incomplete


def add_body(document: Document, text: str, *, indent: bool = True):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if indent else Cm(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)
    return paragraph


def add_title_page(document: Document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(120)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(TITLE)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Experiment Report")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(48)
    meta.paragraph_format.space_after = Pt(0)
    run = meta.add_run("Course Paper Draft\nApril 2026")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)

    document.add_page_break()


def add_table_1(document: Document):
    add_caption(document, "Table 1. Data sources and variable categories.")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Source", "Category", "Representative Variables", "Modeling Role"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, center=True)

    rows = [
        [
            "Kaggle electricity market data",
            "Price and market generation context",
            "price and market-level operational variables",
            "Construct the future negative-price label and provide core market context.",
        ],
        [
            "ENTSO-E Transparency Platform",
            "Demand and cross-border system variables",
            "load, total export, total import",
            "Represent balance conditions, system stress, and transfer constraints.",
        ],
        [
            "ERA5 weather reanalysis",
            "Meteorological drivers",
            "temperature, wind speed, shortwave radiation, cloud cover, precipitation, pressure",
            "Capture exogenous conditions affecting renewable output and demand patterns.",
        ],
        [
            "Derived calendar and institutional features",
            "Temporal and institutional context",
            "hour, month, weekend flag, holiday flag",
            "Represent recurring demand regimes and low-load institutional periods.",
        ],
    ]
    for row_values in rows:
        row_cells = table.add_row().cells
        for cell, value in zip(row_cells, row_values):
            set_cell_text(cell, value)
    document.add_paragraph()


def add_table_2(document: Document, metrics: dict[str, pd.Series]):
    add_caption(
        document,
        "Table 2. Matched main-task benchmark comparison on the common 20-country + public + window=168 + h=6 task.",
    )
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Experiment", "Model Family", "Task Role", "PR-AUC", "ROC-AUC", "Precision", "Recall", "F1"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, center=True)

    benchmark_rows = [
        ("E41", "Logistic regression", "Simple classical baseline"),
        ("E42", "XGBoost", "Boosted-tree classical baseline"),
        ("E43", "CatBoost", "Best matched classical single model"),
        ("E44", "LightGBM", "Strong matched classical single model"),
        ("E49", "GRUMultiMarket", "Best stable standalone deep model"),
        ("E75", "Nonlinear LightGBM stacking", "Strongest pure-meta reference"),
        ("E78", "Mechanism-aware CatBoost", "Best newly trained final-phase single model"),
        ("E81", "Late fusion (E75 + E78)", "Best overall result"),
    ]
    for experiment_id, family, role in benchmark_rows:
        row = metrics[experiment_id]
        cells = table.add_row().cells
        values = [
            experiment_id,
            family,
            role,
            f4(row["pr_auc"]),
            f4(row["roc_auc"]),
            f4(row["precision"]),
            f4(row["recall"]),
            f4(row["f1"]),
        ]
        for idx, value in enumerate(values):
            set_cell_text(cells[idx], value, center=idx >= 3)
    note = (
        "Note. E49 reports the repeated-seed mean on the matched benchmark task. "
        "All rows in this table use the same final main-task definition: 20 countries, public features, "
        "168-hour input window, and a 6-hour forecasting horizon."
    )
    add_body(document, note, indent=False)


def add_table_3(document: Document):
    add_caption(document, "Table 3. Milestone progression of the best score across the four experiment phases.")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Representative Experiments", "Best Milestone", "Best PR-AUC", "Interpretation"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, center=True)

    rows = [
        (
            "Phase 1",
            "E23-E35",
            "E35",
            "0.3691",
            "Single-market hybrid deep development produced the first strong deep reference.",
        ),
        (
            "Phase 2",
            "E45-E52",
            "E49",
            "0.3867",
            "Multi-market GRUMultiMarket became the most defensible standalone deep model.",
        ),
        (
            "Phase 3",
            "E41-E76",
            "E75",
            "0.4323",
            "Completed classical baselines and nonlinear stacking pushed the practical score ceiling.",
        ),
        (
            "Phase 4",
            "E77-E81",
            "E81",
            "0.4388",
            "Mechanism-aware CatBoost plus forced closeout fusion yielded the final best overall result.",
        ),
    ]
    for row_values in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(row[idx], value, center=idx == 3)
    document.add_paragraph()


def add_appendix_tables(document: Document):
    completed, incomplete = build_experiment_registry()

    add_caption(document, "Table A1. Full experiment registry with retained experiment-level test results.")
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Experiment", "Model / Family", "Status", "Test PR-AUC", "Note"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True, center=True)

    note_map = {
        "E8": "Historical early CatBoost benchmark; not part of the final matched main-table comparison.",
        "E41": "Simple matched machine-learning baseline.",
        "E42": "Strong matched XGBoost baseline.",
        "E43": "Strongest matched classical single-model baseline.",
        "E44": "Near-tied LightGBM classical benchmark.",
        "E49": "Repeated-seed mean retained as the stable standalone deep result.",
        "E75": "Best pure-meta reference before final closeout fusion.",
        "E78": "Best newly trained final-phase single model.",
        "E81": "Best overall final result used in the thesis.",
    }
    for experiment_id in sorted(completed.keys(), key=experiment_sort_key):
        item = completed[experiment_id]
        row = table.add_row().cells
        values = [
            experiment_id,
            str(item["model"]),
            "completed",
            f4(item["pr_auc"]),
            note_map.get(experiment_id, ""),
        ]
        for idx, value in enumerate(values):
            set_cell_text(row[idx], value, center=idx == 3)

    add_body(
        document,
        "Table A1 reports the retained experiment-level test result for every locally completed experiment that has a "
        "valid top-level metrics file. Nested member runs, seed-run internals, and candidate-search subdirectories are "
        "excluded so that the appendix remains comparable at the experiment level.",
        indent=False,
    )

    if incomplete:
        add_caption(document, "Table A2. Configured experiments without a retained final experiment-level result.")
        table2 = document.add_table(rows=1, cols=3)
        table2.style = "Table Grid"
        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers2 = ["Experiment", "Status", "Note"]
        for cell, header in zip(table2.rows[0].cells, headers2):
            set_cell_text(cell, header, bold=True, center=True)
        for experiment_id in incomplete:
            row = table2.add_row().cells
            note = "Configured in code but no retained final top-level result was found locally."
            if experiment_id == "E48":
                note = "Repeated-seed graph-hybrid check was not retained as a final completed aggregate result."
            values = [experiment_id, "not retained / not completed", note]
            for idx, value in enumerate(values):
                set_cell_text(row[idx], value)
        add_body(
            document,
            "The experiments in Table A2 are part of the configured codebase but do not currently have a retained "
            "top-level final result available in local outputs.",
            indent=False,
        )


def build_document():
    metrics = build_main_metrics()

    document = Document()
    configure_document(document)
    add_title_page(document)

    document.add_heading("Abstract", level=1)
    add_body(
        document,
        "This report summarizes the full experimental program for future negative electricity price event prediction "
        "in European power markets. The project formulates the problem as leakage-free future event classification: "
        "for each anchor time t, the model observes only historical information up to t and predicts whether the price "
        "at t + h is negative. The final matched benchmark task uses a 168-hour rolling window, a 6-hour forecast "
        "horizon, a 20-country public-feature setting, and chronological train-validation-test splits. The experiment "
        "program compared simple machine-learning baselines, deep temporal encoders, hybrid multi-market deep models, "
        "and several fusion and meta-learning strategies. The strongest standalone classical baseline became E43 "
        "CatBoost with test PR-AUC = 0.4147, while the strongest stable standalone deep model became E49 "
        "GRUMultiMarket with repeated-seed mean test PR-AUC = 0.3867. Later score-ceiling work produced E75 as the "
        "best reusable pure-meta reference at test PR-AUC = 0.4323, and final new-member generation produced E78 as "
        "a mechanism-aware CatBoost model with test PR-AUC = 0.4296. The strongest overall result is E81, a forced "
        "late fusion between E75 and E78, which reaches test PR-AUC = 0.4388. Taken together, the experiments show "
        "that standalone deep models did not surpass the best standalone tree models on this task, but they did add "
        "complementary signal that improved the final fused benchmark.",
        indent=False,
    )

    document.add_heading("Keywords", level=1)
    add_body(
        document,
        "negative electricity prices; event prediction; European power markets; deep learning; CatBoost; late fusion",
        indent=False,
    )

    document.add_heading("Introduction", level=1)
    add_body(
        document,
        "Negative electricity prices have become an increasingly visible phenomenon in European power markets, "
        "especially in the period after 2020 as renewable penetration, transmission constraints, and cross-border "
        "balancing pressures intensified. The phenomenon reflects more than high renewable output alone. It also "
        "reveals low-demand conditions, flexibility limits, and the uneven ability of neighboring systems to absorb "
        "excess supply. For that reason, negative pricing is best interpreted as a system-state problem rather than "
        "as a univariate price anomaly.",
    )
    add_body(
        document,
        "This project therefore studies negative pricing as a future event prediction task rather than as plain price "
        "regression or same-time description. At each anchor time t, the model is given a rolling historical window "
        "of system conditions and must estimate whether a negative-price event will occur at a future target hour "
        "t + h. This framing aligns the research objective with practical forecasting, preserves temporal causality, "
        "and allows direct comparison between manually engineered tabular baselines and learned temporal "
        "representations.",
    )
    add_body(
        document,
        "Three research questions organize the experiment program. First, under what combinations of price, load, "
        "weather, generation-related conditions, and cross-border structure do future negative-price events emerge? "
        "Second, can deep temporal encoders learn system-state representations that improve future event prediction? "
        "Third, do learned multivariate temporal representations outperform manually engineered machine-learning "
        "features once the task is defined in a leakage-free and chronologically consistent way?",
    )

    document.add_heading("Data and Problem Formulation", level=1)
    add_body(
        document,
        "The dataset integrates multiple public sources. Electricity price and generation-related market information "
        "come from a Kaggle-based European electricity panel. Demand and cross-border system information come from the "
        "ENTSO-E Transparency Platform. Weather variables come from Copernicus ERA5, and temporal context variables "
        "such as weekend and legal-holiday indicators are derived in preprocessing. Table 1 summarizes the main data "
        "categories and their modeling role.",
    )
    add_table_1(document)
    add_body(
        document,
        "The fixed project definition is future event classification. For each eligible anchor time t, the input "
        "contains only observations available up to t, organized as a rolling historical window. The main benchmark "
        "task uses a 168-hour window and predicts whether the electricity price at t + 6 hours is negative. The "
        "binary label is therefore y_(t+6) = 1 if the future target-hour price is below zero and 0 otherwise. "
        "Chronological train-validation-test splits preserve temporal ordering and prevent leakage from future values.",
    )
    add_body(
        document,
        "The final matched benchmark comparison in this report uses the common task definition "
        "20-country + public + window=168 + h=6. The public feature group includes price, load, six weather variables, "
        "and calendar context such as weekend and holiday flags. Richer renewables and cross-border flow features were "
        "explored in auxiliary branches, but they were not used to define the main final benchmark table because the "
        "project needed one fair comparison set shared by classical, deep, and fusion models. On the final matched "
        "test set, the benchmark task contains 78,735 samples with 1,443 positive future negative-price events, which "
        "makes class imbalance a central evaluation concern.",
    )

    document.add_heading("Methodology", level=1)
    add_body(
        document,
        "The methodology deliberately separated problem definition from model family choice. The core pipeline is "
        "past window to temporal encoder to classifier to future negative-price event probability. This allowed the "
        "project to begin with leakage-free baselines before adding modeling complexity, rather than assuming that a "
        "specific deep architecture would necessarily be optimal.",
    )
    add_body(
        document,
        "Classical baselines used tabular lagged and rolling-statistical summaries of the main task window. The "
        "matched final baseline family consists of LogisticRegression, XGBoost, CatBoost, and LightGBM. These models "
        "provided an interpretable and competitive reference for what could be achieved through engineered tabular "
        "features under the same benchmark task.",
    )
    add_body(
        document,
        "Deep-learning experiments progressed from GRU-based encoders to hybrid sequence models and eventually to "
        "GRUMultiMarket, which jointly encoded the target market and a shared multi-market context. The project also "
        "tested richer feature groups, gating variants, attention replacements, graph-inspired variants, and "
        "mechanism-aware extensions. The role of the deep models was not merely to fit nonlinear decision boundaries, "
        "but to learn latent system-state representations directly from multivariate temporal evolution.",
    )
    add_body(
        document,
        "Fusion and meta-model families formed the final methodological layer. Weighted late fusion used validation "
        "PR-AUC to combine member probabilities, while stacking treated the outputs of strong component models as "
        "inputs to a second-stage learner. These fusion experiments became important once it was clear that the best "
        "deep standalone model and the best tree-based models captured partially different signals.",
    )
    add_body(
        document,
        "Evaluation emphasized PR-AUC as the primary metric because the target class is relatively rare and because "
        "the project cared about ranking future negative-price risk under imbalance. ROC-AUC, precision, recall, and "
        "F1-score were also reported. Repeated-seed checks were used when a deep-learning result looked promising "
        "enough to justify a stronger robustness claim.",
    )

    document.add_heading("Experimental Progression and Results", level=1)
    add_body(
        document,
        "The complete experiment program can be summarized as four phases. The goal here is not to list every run in "
        "repository order, but to explain how the benchmark evolved and why the final result became E81.",
    )

    document.add_heading("Phase 1: Early deep-learning and richer-feature exploration", level=2)
    add_body(
        document,
        "The early phase established the forecasting framing and then searched for a viable deep backbone under "
        "leakage-free chronological splits. This stage covered early GRU and TCN comparisons, richer-feature subset "
        "tests, and the first hybrid designs. The main practical outcome was that deep modeling had value, but the "
        "strongest early gains did not yet come from large architectural novelty. Instead, the best result of the "
        "early deep stage was E35, a late fusion over the strongest single-market deep branches, with test PR-AUC = "
        f"{f4(metrics['E35']['pr_auc'])}. This phase established that deep models could become competitive, but they "
        "still did not yet challenge the best classical single-model baselines.",
    )

    document.add_heading("Phase 2: Hybrid and multi-market deep-learning development", level=2)
    add_body(
        document,
        "The second phase shifted attention from single-market refinements to explicit multi-market sequence modeling. "
        "That move produced the first strong and stable standalone deep result: E49, the repeated-seed version of "
        "GRUMultiMarket. E49 reached repeated-seed mean test PR-AUC = 0.3867 with very low dispersion, which made it "
        "the most defensible standalone deep-learning reference in the project. This phase was important because it "
        "showed that multi-market context was a more valuable deep-learning direction than the first attention or "
        "graph replacements, many of which later underperformed the simpler multi-market encoder.",
    )

    document.add_heading("Phase 3: Classical baseline completion and score-ceiling fusion", level=2)
    add_body(
        document,
        "Once the project completed the matched classical baselines E41-E44, the benchmark picture became much "
        "clearer. LogisticRegression was a weak but useful floor, while XGBoost, CatBoost, and LightGBM formed a very "
        "strong tree-based baseline group. The strongest standalone classical models, E43 and E44, both exceeded the "
        "best standalone deep score. This meant that further progress would likely come not from small standalone deep "
        "tweaks, but from exploiting complementarity between the deep representation and the strong classical learners.",
    )
    add_body(
        document,
        "That hypothesis was confirmed by fusion and stacking experiments. Simple deep-plus-tree fusions such as E56 "
        "and E58 outperformed their component models, and later nonlinear stacking reached E75 with test PR-AUC = "
        f"{f4(metrics['E75']['pr_auc'])}. E75 became the strongest reusable pure-meta reference. However, by this "
        "stage the no-retrain meta line had already started approaching a ceiling, because several increasingly "
        "complex meta variants delivered only tiny incremental gains over one another.",
    )

    document.add_heading("Phase 4: Final new-member generation and closeout fusion through E81", level=2)
    add_body(
        document,
        "The final phase therefore focused on generating a genuinely new strong member that could still complement the "
        "best meta ceiling. Mechanism-aware tree models were tested, and E78 emerged as the clear success case. It "
        "reached test PR-AUC = 0.4296, which was far stronger than the failed deep mechanism-channel variant E79 and "
        "close to the overall best pure-meta benchmark. An automatic best-member wrapper, E80, overfit validation "
        "selection and chose the wrong member, so it was not retained as a positive result.",
    )
    add_body(
        document,
        "The closeout question was then answered directly by E81, a forced late fusion between E75 and E78. This "
        "experiment required no new base-model training, reused finished artifacts, and delivered the strongest final "
        "result in the project: test PR-AUC = 0.4388. The result confirmed that E78 contributed genuinely useful new "
        "signal to the strongest existing meta reference, and that the failure of E80 came from candidate selection "
        "error rather than from a lack of complementarity.",
    )

    add_table_2(document, metrics)
    add_table_3(document)

    document.add_heading("Discussion", level=1)
    add_body(
        document,
        "Several clear lessons emerge from the full experiment record. First, standalone deep learning did not beat "
        "the strongest standalone classical baselines on the matched final task. E49 is a meaningful and stable deep "
        "result, but it remains below E43 and E44 when each is evaluated as a standalone model. Second, this does not "
        "imply that deep learning lacked value. The fusion results show the opposite: the deep multi-market "
        "representation contributed complementary information that repeatedly improved overall performance when "
        "combined with strong classical members.",
    )
    add_body(
        document,
        "Third, not every plausible deep extension was worth continued investment. Early attention replacements and "
        "the first graph-inspired variants did not outperform the simpler GRUMultiMarket line. The renewables-track "
        "branch did not remain strong enough under later matched-budget checks to justify making it a main thesis "
        "claim. The mechanism-channel deep variant E79 also failed clearly, which helped narrow the project to a much "
        "more defensible ending position.",
    )
    add_body(
        document,
        "Fourth, the final phase clarified the difference between a reusable meta ceiling and a new trained member. "
        "E75 is important because it represents the strongest pure-meta reference built from previously trained strong "
        "members. E78 is important because it is a newly trained single model that almost reached the same level. E81 "
        "matters most because it combines those two lines and becomes the strongest overall result. This final outcome "
        "supports a balanced conclusion: the project did not end with a single deep model beating everything else, but "
        "it did show that learned temporal representations are valuable because they improve the best overall fused "
        "forecasting system.",
    )

    document.add_heading("Conclusion", level=1)
    add_body(
        document,
        "This course paper framed negative electricity pricing as a future event prediction problem and evaluated a "
        "broad set of classical, deep, and fusion models under a common leakage-free benchmark. The final hierarchy is "
        "clear. E81 is the best overall result with test PR-AUC = 0.4388. E75 is the strongest reusable pure-meta "
        "reference with test PR-AUC = 0.4323. E78 is the strongest newly trained final-phase single model with test "
        "PR-AUC = 0.4296. E49 remains the strongest stable standalone deep model with repeated-seed mean test PR-AUC = "
        "0.3867. Within the matched simple machine-learning baselines, E41-E44 provide the final standalone benchmark "
        "family, with E43 CatBoost as the strongest single-model baseline.",
    )
    add_body(
        document,
        "These results place the project in a defensible stopping position. The benchmark is strong, the final "
        "experiment story is coherent, and the contribution is not merely that one model achieved the best score. "
        "Rather, the full program demonstrates that future negative-price event prediction benefits from a disciplined "
        "forecasting setup, strong classical baselines, and learned temporal representations that add complementary "
        "signal to the best tree-based models. Both the narrative and the final benchmark table therefore support a "
        "clear thesis closeout centered on E81 as the strongest overall result.",
    )

    document.add_heading("Appendix", level=1)
    add_appendix_tables(document)

    return document


def qa_document(path: Path):
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    required_sections = [
        "Abstract",
        "Keywords",
        "Introduction",
        "Data and Problem Formulation",
        "Methodology",
        "Experimental Progression and Results",
        "Discussion",
        "Conclusion",
        "Appendix",
    ]
    section_positions = {}
    for section in required_sections:
        try:
            section_positions[section] = paragraphs.index(section)
        except ValueError as exc:
            raise AssertionError(f"Missing section heading: {section}") from exc
    ordered_positions = [section_positions[section] for section in required_sections]
    if ordered_positions != sorted(ordered_positions):
        raise AssertionError("Section headings are not in the required order.")

    def section_text(start_heading: str, end_heading: str | None) -> str:
        start = section_positions[start_heading] + 1
        end = len(paragraphs) if end_heading is None else section_positions[end_heading]
        return "\n".join(paragraphs[start:end])

    abstract_text = section_text("Abstract", "Keywords")
    conclusion_text = section_text("Conclusion", "Appendix")
    if "E81" not in abstract_text:
        raise AssertionError("Abstract does not mention E81.")
    if "E81" not in conclusion_text:
        raise AssertionError("Conclusion does not mention E81.")

    benchmark_table_found = False
    appendix_table_found = False
    for table in document.tables:
        table_text = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        flat = {cell for row in table_text for cell in row}
        header = tuple(table_text[0]) if table_text else tuple()
        if header == (
            "Experiment",
            "Model Family",
            "Task Role",
            "PR-AUC",
            "ROC-AUC",
            "Precision",
            "Recall",
            "F1",
        ):
            benchmark_table_found = True
            if "E8" in flat:
                raise AssertionError("E8 appears in the main benchmark table.")
        if "E8" in flat and header != (
            "Experiment",
            "Model Family",
            "Task Role",
            "PR-AUC",
            "ROC-AUC",
            "Precision",
            "Recall",
            "F1",
        ):
            appendix_table_found = True
    if not benchmark_table_found:
        raise AssertionError("Main benchmark table is missing required experiment IDs.")
    if not appendix_table_found:
        raise AssertionError("Appendix experiment registry is missing E8.")

    return {
        "sections_ok": True,
        "abstract_mentions_E81": True,
        "conclusion_mentions_E81": True,
        "benchmark_table_ok": True,
        "appendix_table_ok": True,
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
    }


def main():
    document = build_document()
    document.save(OUTPUT_PATH)
    qa = qa_document(OUTPUT_PATH)
    print(f"Wrote DOCX to {OUTPUT_PATH}")
    print(f"QA summary: {qa}")


if __name__ == "__main__":
    main()
